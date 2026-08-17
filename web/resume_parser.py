"""简历文件文本提取层。"""

from pathlib import Path


def _read_text_with_fallback(path):
    raw = path.read_bytes()
    for encoding in ("utf-8", "gbk", "gb18030"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _dedupe_repeated_lines(lines):
    seen_first_lines = {}
    result = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped in seen_first_lines and len(result) > 3:
            continue
        if stripped not in seen_first_lines:
            seen_first_lines[stripped] = True
        result.append(line)
    return result


def extract_resume_text(file_path):
    """按扩展名提取简历纯文本。

    返回 {"text": ..., "pages": 页数或 None, "warning": 警告或 None}
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext in (".txt", ".md"):
        text = _read_text_with_fallback(path)
        pages = None
    elif ext == ".pdf":
        try:
            import pdfplumber
        except ImportError:
            raise ValueError("缺少 pdfplumber 依赖，无法解析 PDF")
        pages = 0
        parts = []
        with pdfplumber.open(path) as pdf:
            pages = len(pdf.pages)
            for page in pdf.pages:
                parts.append(page.extract_text() or "")
        text = "\n".join(parts)
    elif ext == ".docx":
        try:
            import docx
        except ImportError:
            raise ValueError("缺少 python-docx 依赖，无法解析 DOCX")
        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
        pages = None
    else:
        raise ValueError("不支持的文件类型")

    lines = _dedupe_repeated_lines(text.splitlines())
    text = "\n".join(lines).strip()
    warning = None
    if len(text) < 100:
        warning = "疑似扫描件或空文件，请改用文字版简历"
    return {"text": text, "pages": pages, "warning": warning}
